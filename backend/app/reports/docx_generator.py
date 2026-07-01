"""Generates templated vendor performance Word reports from an aggregated
weekly summary DataFrame, with conditional highlighting on risk rows.
"""

import os

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.environ.get("REPORT_OUTPUT_DIR", "data/reports")


def _set_cell_shading(cell, hex_color: str):
    """Applies background shading to a table cell (python-docx has no direct API for this)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def build_vendor_report_doc(vendor_id: int, summary_df: pd.DataFrame) -> str:
    doc = Document()

    title = doc.add_heading(f"Vendor Performance Report — Vendor #{vendor_id}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Automated weekly analytics summary.").italic = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Week", "Orders", "Units Sold", "Revenue (INR)", "Failed Orders"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["week"])
        cells[1].text = str(int(row["orders"]))
        cells[2].text = str(int(row["units_sold"]))
        cells[3].text = f"{row['revenue_inr']:.2f}"
        cells[4].text = str(int(row["failed_orders"]))

        failure_rate = row["failed_orders"] / row["orders"] if row["orders"] else 0
        if failure_rate >= 0.08:
            _set_cell_shading(cells[4], "F4CCCC")  # red flag
        elif failure_rate >= 0.02:
            _set_cell_shading(cells[4], "FFF2CC")  # amber flag

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated automatically — no manual edits required.")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, f"vendor_{vendor_id}_report.docx")
    doc.save(out_path)
    return out_path
